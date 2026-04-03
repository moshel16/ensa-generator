"""
Vercel Serverless Function — POST /api/generate-docx
מקבל: { summary: str, topic: str }
מחזיר: קובץ .docx בסגנון ACL_Final — גישת template
"""
from http.server import BaseHTTPRequestHandler
import json, io, re, os, base64

# נתיב ל-template ב-Vercel
TEMPLATE_PATH = os.path.join(os.path.dirname(__file__), '..', 'ACL_Final.docx')

def build_docx(summary: str, topic: str) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from datetime import date

    # ===== פלטת צבעים =====
    NAVY      = RGBColor(0x1B, 0x3A, 0x6B)
    GOLD      = RGBColor(0xD4, 0xA8, 0x00)
    WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
    LIGHT_BLUE= RGBColor(0xD6, 0xE8, 0xF7)
    BLUE      = RGBColor(0x2E, 0x5F, 0xA3)
    DARK_TEXT = RGBColor(0x2C, 0x3E, 0x50)
    GREEN     = RGBColor(0x1E, 0x7A, 0x3C)
    TEAL      = RGBColor(0x0F, 0x7B, 0x6C)
    CODE_TEXT = RGBColor(0xA8, 0xFF, 0x78)
    SUBTEXT   = RGBColor(0xAA, 0xAA, 0xAA)
    GRAY_TEXT = RGBColor(0x5D, 0x6D, 0x7E)

    NAVY_HEX  = "1B3A6B"
    BLUE_HEX  = "2E5FA3"
    LB_HEX    = "D6E8F7"
    GOLD_HEX  = "D4A800"
    TEAL_HEX  = "0F7B6C"
    WHITE_HEX = "FFFFFF"
    CODE_BG   = "1C1C1C"
    GREEN_BG  = "D4EDDA"

    def hex2rgb(h):
        h = h.lstrip('#')
        return RGBColor(int(h[0:2],16), int(h[2:4],16), int(h[4:6],16))

    # ===== טען template =====
    doc = Document(TEMPLATE_PATH)

    # ===== מחק את כל תוכן הגוף, שמור sectPr =====
    body = doc.element.body
    sectPr = body.find(qn('w:sectPr'))
    for child in list(body):
        body.remove(child)
    if sectPr is not None:
        body.append(sectPr)

    # ===== helpers =====
    def set_cell_bg(cell, hex_color):
        tcPr = cell._tc.get_or_add_tcPr()
        # הסר shd קיים
        for old in tcPr.findall(qn('w:shd')):
            tcPr.remove(old)
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), hex_color)
        tcPr.append(shd)

    def set_cell_borders(cell, color_hex, size=6):
        tcPr = cell._tc.get_or_add_tcPr()
        for old in tcPr.findall(qn('w:tcBorders')):
            tcPr.remove(old)
        tcBorders = OxmlElement('w:tcBorders')
        for side in ['top','left','bottom','right']:
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:val'), 'single')
            el.set(qn('w:sz'), str(size))
            el.set(qn('w:space'), '0')
            el.set(qn('w:color'), color_hex)
            tcBorders.append(el)
        tcPr.append(tcBorders)

    def set_para_shd(para, fill_hex):
        pPr = para._p.get_or_add_pPr()
        for old in pPr.findall(qn('w:shd')):
            pPr.remove(old)
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), fill_hex)
        pPr.append(shd)

    def add_border_bottom(para, color_hex, size=10):
        pPr = para._p.get_or_add_pPr()
        pBdr = pPr.find(qn('w:pBdr'))
        if pBdr is None:
            pBdr = OxmlElement('w:pBdr')
            pPr.insert(0, pBdr)
        bot = OxmlElement('w:bottom')
        bot.set(qn('w:val'), 'single')
        bot.set(qn('w:sz'), str(size))
        bot.set(qn('w:space'), '1')
        bot.set(qn('w:color'), color_hex)
        pBdr.append(bot)

    def add_border_top(para, color_hex, size=10):
        pPr = para._p.get_or_add_pPr()
        pBdr = pPr.find(qn('w:pBdr'))
        if pBdr is None:
            pBdr = OxmlElement('w:pBdr')
            pPr.insert(0, pBdr)
        top = OxmlElement('w:top')
        top.set(qn('w:val'), 'single')
        top.set(qn('w:sz'), str(size))
        top.set(qn('w:space'), '1')
        top.set(qn('w:color'), color_hex)
        pBdr.insert(0, top)

    def set_spacing(para, before=0, after=0):
        pPr = para._p.get_or_add_pPr()
        for old in pPr.findall(qn('w:spacing')):
            pPr.remove(old)
        spc = OxmlElement('w:spacing')
        if before: spc.set(qn('w:before'), str(before))
        if after:  spc.set(qn('w:after'),  str(after))
        pPr.append(spc)

    def run_font(run, size_pt, color, bold=False, font='Arial', italic=False, no_rtl=False):
        run.font.name = font
        run.font.size = Pt(size_pt)
        run.font.color.rgb = color
        run.bold = bold
        run.italic = italic
        rPr = run._r.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.insert(0, rFonts)
        rFonts.set(qn('w:ascii'), font)
        rFonts.set(qn('w:hAnsi'), font)
        if not no_rtl:
            # RTL מפורש — מבטיח יישור ימין בשורות ארוכות
            if rPr.find(qn('w:rtl')) is None:
                rtl_el = OxmlElement('w:rtl')
                rPr.append(rtl_el)
        if bold:
            if rPr.find(qn('w:bCs')) is None:
                bCs = OxmlElement('w:bCs')
                rPr.append(bCs)
        if italic:
            if rPr.find(qn('w:iCs')) is None:
                iCs = OxmlElement('w:iCs')
                rPr.append(iCs)

    def make_content_table(doc, fill_fn, border_hex=BLUE_HEX, fill_hex=LB_HEX):
        """
        בונה טבלה חד-עמודה עם bidiVisual — זהה ל-ACL_Final.
        fill_fn(cell) ממלאת את התא.
        """
        table = doc.add_table(rows=1, cols=1)

        tbl = table._tbl
        existing = tbl.find(qn('w:tblPr'))
        if existing is not None:
            tbl.remove(existing)

        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)

        tblPr.append(OxmlElement('w:bidiVisual'))

        tblW = OxmlElement('w:tblW')
        tblW.set(qn('w:w'), '9200')
        tblW.set(qn('w:type'), 'dxa')
        tblPr.append(tblW)

        tblInd = OxmlElement('w:tblInd')
        tblInd.set(qn('w:w'), '140')
        tblInd.set(qn('w:type'), 'dxa')
        tblPr.append(tblInd)

        tblCellMar = OxmlElement('w:tblCellMar')
        for side, val in [('left','10'),('right','10')]:
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:w'), val)
            el.set(qn('w:type'), 'dxa')
            tblCellMar.append(el)
        tblPr.append(tblCellMar)

        tblLook = OxmlElement('w:tblLook')
        tblLook.set(qn('w:val'), '0000')
        for attr in ['firstRow','lastRow','firstColumn','lastColumn','noHBand','noVBand']:
            tblLook.set(qn(f'w:{attr}'), '0')
        tblPr.append(tblLook)

        cell = table.cell(0, 0)

        tcPr = cell._tc.get_or_add_tcPr()
        tcW = OxmlElement('w:tcW')
        tcW.set(qn('w:w'), '9200')
        tcW.set(qn('w:type'), 'dxa')
        tcPr.insert(0, tcW)

        tcMar = OxmlElement('w:tcMar')
        for side, val in [('top','100'),('left','140'),('bottom','100'),('right','140')]:
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:w'), val)
            el.set(qn('w:type'), 'dxa')
            tcMar.append(el)
        tcPr.append(tcMar)

        set_cell_bg(cell, fill_hex)
        set_cell_borders(cell, border_hex)

        for p in cell.paragraphs:
            p._p.getparent().remove(p._p)

        fill_fn(cell)
        return table

    def cell_add_para(cell, ltr=False):
        """
        מוסיף פסקה לתא.
        בתוך bidiVisual: jc=left = ימין ויזואלי, jc=right = שמאל ויזואלי.
        ltr=True לפקודות שצריכות LTR.
        """
        from docx.text.paragraph import Paragraph as DocxPara
        p = cell._tc.add_p()
        para = DocxPara(p, cell._tc)
        pPr = para._p.get_or_add_pPr()
        if ltr:
            bidi0 = OxmlElement('w:bidi')
            bidi0.set(qn('w:val'), '0')
            pPr.insert(0, bidi0)
            jcEl = OxmlElement('w:jc')
            jcEl.set(qn('w:val'), 'right')
            pPr.append(jcEl)
        else:
            jcEl = OxmlElement('w:jc')
            jcEl.set(qn('w:val'), 'left')
            pPr.append(jcEl)
        return para

    def is_ltr_line(text):
        """זיהוי שורה שהיא פקודה/קוד — צריכה LTR"""
        t = text.strip().lower()
        ltr_keywords = ['show ', 'debug ', 'ip ', 'access-list', 'interface',
                        'router ', 'network ', 'permit', 'deny', 'no ',
                        'exit', 'ping', 'traceroute', 'copy ', 'write',
                        'clear ', 'line vty', 'access-class', 'access-group']
        return any(t.startswith(k) for k in ltr_keywords)

    def parse_inline(para, text, default_color=None):
        """מפרסר **bold** inline"""
        dc = default_color or DARK_TEXT
        parts = re.split(r'(\*\*[^*]+\*\*)', text)
        for part in parts:
            if not part: continue
            if part.startswith('**') and part.endswith('**') and len(part) > 4:
                run = para.add_run(part[2:-2])
                run_font(run, 11, BLUE, bold=True)
            else:
                clean = part.replace('**', '')
                if not clean: continue
                run = para.add_run(clean)
                run_font(run, 11, dc)

    # ===== כותרת ראשית =====
    p = doc.add_paragraph()
    set_para_shd(p, NAVY_HEX)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_border_top(p, GOLD_HEX, 12)
    set_spacing(p, before=160)
    r1 = p.add_run('סיכום ')
    run_font(r1, 26, GOLD, bold=True)
    r2 = p.add_run(topic)
    run_font(r2, 22, GOLD, bold=True)

    p2 = doc.add_paragraph()
    set_para_shd(p2, NAVY_HEX)
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p2.add_run('CCNA: Enterprise Networking, Security & Automation')
    run_font(r, 12, LIGHT_BLUE)

    p3 = doc.add_paragraph()
    set_para_shd(p3, NAVY_HEX)
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_border_bottom(p3, GOLD_HEX, 12)
    set_spacing(p3, after=140)
    today = date.today().strftime('%d.%m.%Y')
    for text, color, bold in [
        ('הכין: משה לופו', hex2rgb(GREEN_BG), True),
        ('   |   ',         SUBTEXT,           False),
        ('מרצה: אלי פדידה', WHITE,             True),
        ('   |   ',         SUBTEXT,           False),
        (today,            LIGHT_BLUE,         False),
    ]:
        r = p3.add_run(text)
        run_font(r, 14, color, bold=bold)

    doc.add_paragraph()

    # ===== עיבוד תוכן =====
    lines = summary.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i]

        # --- בלוק קוד ---
        if line.startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].startswith('```'):
                code_lines.append(lines[i])
                i += 1
            if code_lines:
                doc.add_paragraph()
                for cl in code_lines:
                    cp = doc.add_paragraph()
                    pPr = cp._p.get_or_add_pPr()
                    # LTR מוחלט לקוד
                    bidi0 = OxmlElement('w:bidi')
                    bidi0.set(qn('w:val'), '0')
                    pPr.insert(0, bidi0)
                    jcEl = OxmlElement('w:jc')
                    jcEl.set(qn('w:val'), 'left')
                    pPr.append(jcEl)
                    set_para_shd(cp, CODE_BG)
                    set_spacing(cp, before=0, after=0)
                    cr = cp.add_run(cl if cl.strip() else ' ')
                    run_font(cr, 9, CODE_TEXT, font='Courier New')
                    rPr = cr._r.get_or_add_rPr()
                    rtl0 = OxmlElement('w:rtl')
                    rtl0.set(qn('w:val'), '0')
                    rPr.append(rtl0)
                doc.add_paragraph()
            i += 1
            continue

        # --- כותרת H1 ---
        if line.startswith('# ') and not line.startswith('## '):
            doc.add_paragraph()
            p = doc.add_paragraph()
            set_para_shd(p, NAVY_HEX)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_border_bottom(p, GOLD_HEX, 10)
            set_spacing(p, before=220, after=130)
            r = p.add_run(line[2:])
            run_font(r, 15, WHITE, bold=True, no_rtl=True)
            i += 1; continue

        # --- כותרת H2 ---
        if line.startswith('## ') and not line.startswith('### '):
            doc.add_paragraph()
            p = doc.add_paragraph()
            set_para_shd(p, NAVY_HEX)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_border_bottom(p, GOLD_HEX, 10)
            set_spacing(p, before=220, after=130)
            r = p.add_run(line[3:])
            run_font(r, 15, WHITE, bold=True, no_rtl=True)
            i += 1; continue

        # --- כותרת H3 ---
        if line.startswith('### '):
            h3_text = line[4:]
            def fill_h3(cell, t=h3_text):
                para = cell_add_para(cell)
                set_spacing(para, before=100, after=60)
                add_border_bottom(para, BLUE_HEX, 4)
                r = para.add_run(t)
                run_font(r, 12, NAVY, bold=True)
            make_content_table(doc, fill_h3, BLUE_HEX, LB_HEX)
            i += 1; continue

        # --- שורת הסבר ---
        if re.match(r'^הסבר\s*:', line):
            explain = re.sub(r'^הסבר\s*:\s*', '', line)
            def fill_explain(cell, t=explain):
                para = cell_add_para(cell)
                set_spacing(para, before=40, after=10)
                r2 = para.add_run(t)
                run_font(r2, 11, DARK_TEXT, italic=True)
            make_content_table(doc, fill_explain, TEAL_HEX, 'D4F0EB')
            i += 1; continue

        # --- רשימה ממוספרת ---
        if re.match(r'^\d+\.\s', line):
            items = []
            while i < len(lines) and re.match(r'^\d+\.\s', lines[i]):
                items.append(lines[i])
                i += 1
            def fill_numbered(cell, its=items):
                for item in its:
                    m = re.match(r'^(\d+\.)\s+(.*)', item)
                    if not m: continue
                    num, rest = m.group(1), m.group(2)
                    para = cell_add_para(cell)
                    set_spacing(para, before=35, after=35)
                    rn = para.add_run(num + ' ')
                    run_font(rn, 11, BLUE, bold=True)
                    parse_inline(para, rest)
            make_content_table(doc, fill_numbered)
            continue

        # --- bullets ---
        if line.startswith('- ') or line.startswith('* '):
            bullets = []
            while i < len(lines) and (lines[i].startswith('- ') or lines[i].startswith('* ')):
                bullets.append(lines[i][2:])
                i += 1
            def fill_bullets(cell, bs=bullets):
                for b in bs:
                    para = cell_add_para(cell)
                    set_spacing(para, before=30, after=30)
                    rd = para.add_run('◆ ')
                    run_font(rd, 11, BLUE, bold=True)
                    parse_inline(para, b)
            make_content_table(doc, fill_bullets)
            continue

        # --- טבלת Markdown ---
        if line.startswith('|'):
            tbl_lines = []
            while i < len(lines) and lines[i].startswith('|'):
                tbl_lines.append(lines[i])
                i += 1
            rows_data = [l for l in tbl_lines if not re.match(r'^\|[\s:-]+\|', l)]
            if not rows_data:
                continue
            col_count = len(rows_data[0].split('|')) - 2
            if col_count < 1: continue

            TW = 9200
            if col_count == 2:   cws = [2760, 6440]
            elif col_count == 3: cws = [2300, 3450, 3450]
            else:                cws = [TW // col_count] * col_count

            table = doc.add_table(rows=len(rows_data), cols=col_count)

            tbl = table._tbl
            existing = tbl.find(qn('w:tblPr'))
            if existing is not None:
                tbl.remove(existing)
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)
            tblPr.append(OxmlElement('w:bidiVisual'))
            tblW = OxmlElement('w:tblW')
            tblW.set(qn('w:w'), str(TW))
            tblW.set(qn('w:type'), 'dxa')
            tblPr.append(tblW)
            tblInd = OxmlElement('w:tblInd')
            tblInd.set(qn('w:w'), '140')
            tblInd.set(qn('w:type'), 'dxa')
            tblPr.append(tblInd)
            tblLook = OxmlElement('w:tblLook')
            tblLook.set(qn('w:val'), '0000')
            for attr in ['firstRow','lastRow','firstColumn','lastColumn','noHBand','noVBand']:
                tblLook.set(qn(f'w:{attr}'), '0')
            tblPr.append(tblLook)

            for ri, row_line in enumerate(rows_data):
                cells_text = [c.strip() for c in row_line.split('|')[1:-1]]
                is_header = (ri == 0)
                row = table.rows[ri]
                for ci, cell in enumerate(row.cells[:col_count]):
                    fill = NAVY_HEX if is_header else (LB_HEX if ri % 2 == 0 else WHITE_HEX)
                    set_cell_bg(cell, fill)
                    set_cell_borders(cell, BLUE_HEX)

                    tcPr = cell._tc.get_or_add_tcPr()
                    tcW = OxmlElement('w:tcW')
                    tcW.set(qn('w:w'), str(cws[ci] if ci < len(cws) else TW // col_count))
                    tcW.set(qn('w:type'), 'dxa')
                    tcPr.insert(0, tcW)
                    tcMar = OxmlElement('w:tcMar')
                    for side, val in [('top','80'),('left','120'),('bottom','80'),('right','120')]:
                        el = OxmlElement(f'w:{side}')
                        el.set(qn('w:w'), val)
                        el.set(qn('w:type'), 'dxa')
                        tcMar.append(el)
                    tcPr.append(tcMar)

                    for p in cell.paragraphs:
                        p._p.getparent().remove(p._p)

                    from docx.text.paragraph import Paragraph as DocxPara
                    cp = cell._tc.add_p()
                    para = DocxPara(cp, cell._tc)
                    pPr = para._p.get_or_add_pPr()
                    jcEl = OxmlElement('w:jc')
                    jcEl.set(qn('w:val'), 'center' if is_header else 'right')
                    pPr.append(jcEl)
                    set_spacing(para, before=20, after=20)

                    txt = re.sub(r'`([^`]+)`', r'\1', cells_text[ci] if ci < len(cells_text) else '')
                    r = para.add_run(txt)
                    run_font(r, 10 if not is_header else 11,
                             WHITE if is_header else DARK_TEXT,
                             bold=is_header)

            doc.add_paragraph()
            continue

        # --- שורה ריקה ---
        if not line.strip():
            doc.add_paragraph()
            i += 1; continue

        # --- טקסט רגיל ---
        text = line.strip()
        if text:
            def fill_text(cell, t=text):
                ltr = is_ltr_line(t)
                para = cell_add_para(cell, ltr=ltr)
                set_spacing(para, before=20, after=20)
                if ltr:
                    r = para.add_run(t)
                    run_font(r, 11, DARK_TEXT, font='Courier New')
                    rPr = r._r.get_or_add_rPr()
                    rtl0 = OxmlElement('w:rtl')
                    rtl0.set(qn('w:val'), '0')
                    rPr.append(rtl0)
                else:
                    parse_inline(para, t)
            make_content_table(doc, fill_text)
        i += 1

    # ===== Footer =====
    doc.add_paragraph()
    fp = doc.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_border_top(fp, BLUE_HEX, 4)
    set_spacing(fp, before=100)
    fr = fp.add_run(f'CCNA ENSA | {topic} | הכין: משה לופו | מרצה: אלי פדידה')
    run_font(fr, 9, GRAY_TEXT)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


class handler(BaseHTTPRequestHandler):
    def do_POST(self):
        length = int(self.headers.get('Content-Length', 0))
        body = json.loads(self.rfile.read(length))
        summary = body.get('summary', '')
        topic   = body.get('topic', 'ENSA')

        try:
            docx_bytes = build_docx(summary, topic)
            self.send_response(200)
            self.send_header('Content-Type', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            self.send_header('Content-Disposition', f'attachment; filename="ENSA_{topic}_Summary.docx"')
            self.send_header('Content-Length', str(len(docx_bytes)))
            self.send_header('Access-Control-Allow-Origin', '*')
            self.end_headers()
            self.wfile.write(docx_bytes)
        except Exception as e:
            self.send_response(500)
            self.send_header('Content-Type', 'application/json')
            self.send_header('Access-Control-Allow-Origin', '*')
            self.end_headers()
            self.wfile.write(json.dumps({'error': str(e)}).encode())

    def do_OPTIONS(self):
        self.send_response(200)
        self.send_header('Access-Control-Allow-Origin', '*')
        self.send_header('Access-Control-Allow-Methods', 'POST, OPTIONS')
        self.send_header('Access-Control-Allow-Headers', 'Content-Type')
        self.end_headers()
